#!/usr/bin/python
# -*- coding: utf-8 -*-
import argparse
import os
import glob
import datetime


def read_summary_arm_results_csv_file(filename, sep='\t'):
    if os.path.exists(filename) == True:
        reclist = []
        f = open(filename,'r')
        for n,line in enumerate(f):
            line = line.strip().split(sep)
            if n == 0:
                header = line
            elif n == 1:
                data = line
                dataDict = dict(zip(header, data))
        f.close()

        atbDic= {}
        for key in dataDict.keys():
            keys = key.split('::')
            atb = keys[0]
            res  = keys[1]
            values = dataDict[key].split(',')
            dic = {}
            for value in values:
                key, value = value.split(':')
                dic[key] = value
            try:
                atbDic[atb][res] = dic
            except KeyError:
                atbDic[atb] = {res:dic}

        return atbDic
    else:
        print '\nNo armDB result file %s\n' % filename
        exit(1)


def read_mlst_results_csv_file(filename, sep='\t'):
    if os.path.exists(filename)==True:
        f = open(filename,'r')
        for n,line in enumerate(f):
            line = line.strip().split(sep)
            if n == 0:
                header = line
            elif n == 1:
                dataDic = dict(zip(header,line))
        f.close()
        return dataDic
    else:
        print '\nNo MLST result file %s\n' % filename
        return ''


def read_species(samplefile, sampleID, sep='\t'):
    if os.path.exists(samplefile)==True:
        f = open(samplefile, 'r')
        for n,line in enumerate(f):
            if line.split(sep)[0] == sampleID:
                species = line.strip().split('\t')[1]
                sampleID = line.strip().split('\t')[0]
                break
        f.close()
        return species
    else:
        print '\nNo sample file %s\n' % samplefile
        exit(1)


def write_docx(wkdir, sampleID, species, ST, amrDic, armDBname, initial):
    atbs = amrDic.keys()
    atbs.sort()
    fresDict = {}
    for atb in atbs:
        results = amrDic[atb].keys()
        results.sort()
        for res in results:
            warning = 0
            if 'warning' in amrDic[atb][res].keys():
                warning = 1
            id  = float(amrDic[atb][res]['id'])
            cov = float(amrDic[atb][res]['cov'])
            dep = amrDic[atb][res]['dep']
            try:
                snp = amrDic[atb][res]['snp']
                snp = snp.split('[')[0].replace('|',',')
                var_only = '1'
            except KeyError:
                snp = ''
                var_only = '0'

            if id >= 90 and cov >= 80:
                res = res.replace('_',' ')
                if snp != '' and snp != 'None':
                    snp1 =  'variant résistant de '.decode('utf-8')
                    snp2 = '%s (%s)' % (res.strip().split(' ')[0], snp)
                    res = snp1 + snp2
                elif snp == 'None':
                    res = ''
                elif id < 100 :
                    res = res + '-like'

                if species in ['klebsiella pneumoniae', 'enterobacter cloacae'] and 'oqx' in res:
                    res = ''

                if 'crrA' in res:
                    res = ''

                if res == 'ampC [Escherichia coli]-like' or res == 'ampC [E. coli]-like' \
                        or res == 'ampC [Escherichia coli]' or res == 'ampC [E. coli]'\
                        or res == 'AmpC [Escherichia coli]' or res == 'AmpC [E. coli]-like'\
                        or res == 'AmpC [Escherichia coli]-like':
                    res = 'AmpC [E. coli]'
                
                if warning == 1:
                    res = res + '*'

            if res.strip() != '':
                try:
                    fresDict[atb].append(res)
                except:
                    fresDict[atb] = [res]


    from docx import Document
    from docx.shared import Inches

    document = Document()
    sections = document.sections
    sections[0].left_margin = Inches(0.5)
    sections[0].right_margin = Inches(0.5)

    ref = document.add_heading('Référence CNR de l\'isolat bacterien : \t'.decode('utf-8'), 3)
    ref.add_run(sampleID).italic = True

    spp = document.add_heading('Espèce bactérienne :\t'.decode('utf-8'), 3)
    spp.add_run(species[0].upper() + species[1:] + '\n').italic = True

    document.add_heading('Matériel & Méthodes :'.decode('utf-8'), 3)
    met_p1 = document.add_paragraph('   ~ Séquençage : '.decode('utf-8'))
    met_p1.add_run('Méthode Illumina (2 x 150pb ou 300pb appariés)\n\n'.decode('utf-8'))
    met_p1.add_run('   ~ Analyse '.decode('utf-8'))
    met_p1.add_run('in silico '.decode('utf-8')).italic = True
    met_p1.add_run('des séquences génomiques : Bowtie2, CD-HIT, MUMmer et Samtools\n\n'.decode('utf-8'))
    if ST != '':
        met_p1.add_run('   ~ Bases de données MLST : '.decode('utf-8'))
        if species == 'escherichia coli':
            met_p1.add_run('http://mlst.warwick.ac.uk/mlst\n\n'.decode('utf-8'))
        elif species == 'klebsiella pneumoniae':
            met_p1.add_run('http://bigsdb.pasteur.fr\n\n'.decode('utf-8'))
        else:
            met_p1.add_run('https://pubmlst.org/databases\n\n'.decode('utf-8'))

    met_p1.add_run('   ~ Bases de données du CNR de la résistance aux antibiotiques : %s ver.: %s cat.: %s\n'.decode('utf-8')
                   % (armDBname.split('_')[0], armDBname.split('_')[1], armDBname.split('_')[2]))

    if ST != '':
        document.add_heading('Résultat : Génotypage MLST '.decode('utf-8'), 3)
        document.add_paragraph('   ~ Sequence Type: ST-%s\n'.decode('utf-8') % ST)

    document.add_heading('Résultat : Déterminants de la résistance aux 3 principales familles d\'antibiotiques (*)'.decode('utf-8'), 3)


    frDict = {'Aminoglycoside':"Aminosides",'Beta-lactam':"Beta-lactamines".decode('utf-8'),'Quinolone':"Quinolones",'Colistin':"Colistine",
              'Sulfonamide':"Sulfamides",'Trimethoprime':"Triméthoprime".decode('utf-8'),'Cycline':"Tétracycline".decode('utf-8'),'Aminoglycoside|Fluoroquinolone':"Aminosides et fluoroquinolones"}
    fresults = ['Aminoglycoside','Aminoglycoside|Fluoroquinolone','Beta-lactam','Quinolone','Colistin'] #,Sulfonamide','Trimethoprime','Cycline',]
    for func in fresults:
        res_txt = ''
        try:
            res = fresDict[func]
            res.sort()
            res_txt = res_txt + '   ~ %s : %s\n'.decode('utf-8') % (frDict[func], ', '.join(res))
            #document.add_paragraph(res_txt)
        except KeyError:
            res_txt = res_txt + '   ~ %s :\n'.decode('utf-8') % (frDict[func])
        document.add_paragraph(res_txt)
            #res_txt = ''

    document.add_paragraph('(*) Contacter le CNR pour d\'autres familles d\'antibiotiques (tél: 04 73 754 920)'.decode('utf-8'))
    outfile = os.path.join(os.path.dirname(wkdir), 'CR_CNR_%s_%s_%s.docx' % (sampleID, datetime.date.today(), initial))
    document.save(outfile)

    print 'The main results are summarized in file %s' % outfile


def main(args):
    wkdir = args.workdir
    initial = args.initial
    sampleID    = os.path.basename(wkdir)
    samplefile  = os.path.join(wkdir, 'sample.csv')
    species     = read_species(samplefile, sampleID, sep='\t')
    try:
        STfilename  = glob.glob(os.path.join(wkdir, 'mlst_report.tsv'))[0]
        STdic       = read_mlst_results_csv_file(STfilename, sep='\t')
        ST = STdic['ST']
    except IndexError:
        ST = ''
    armfilename = glob.glob(os.path.join(wkdir, 'summary_results_armDB_*.csv'))[0]
    armDBname   = os.path.splitext(os.path.basename(armfilename).replace('summary_results_',''))[0]
    amrDic      = read_summary_arm_results_csv_file(armfilename, sep='\t')
    write_docx(wkdir, sampleID, species, ST, amrDic, armDBname, initial)


def version():
	return "1.0"


def run():
    parser = argparse.ArgumentParser(description='write_docx.py - Version ' + version())
    parser.add_argument('-wd', '--workdir',   dest="workdir",  default='COL0027', help='working directory')
    parser.add_argument('-in', '--initial', dest="initial",  default='RBO',     help="Initials of the user")
    parser.add_argument('-V',  '--version',  action='version', version='write_docx.py-' + version(), help="Prints version number")
    args = parser.parse_args()
    main(args)

if __name__ == '__main__':
	run()
