#!/usr/bin/python3
# -*- coding: utf-8 -*-
import argparse
import os
import glob
import datetime

from docx import Document
from docx.shared import Inches


def read_summary_arm_results_csv_file(filename, sep='\t'):
    if os.path.exists(filename):
        dataDict = {}
        with open(filename, 'r') as f:
            header = ""
            for n, line in enumerate(f):
                line = line.strip().split(sep)
                if n == 0:
                    header = line
                elif n == 1:
                    data = line
                    dataDict = dict(zip(header, data))

        atbDic = {}
        for key in dataDict.keys():
            keys = key.split('::')
            atb = keys[0]
            res = keys[1]
            values = dataDict[key].split(',')
            dic = {}
            for value in values:
                key, value = value.split(':')
                dic[key] = value
            try:
                atbDic[atb][res] = dic
            except KeyError:
                atbDic[atb] = {res: dic}

        return atbDic
    else:
        print('\nNo armDB result file {0}\n'.format(filename))
        exit(1)


def read_mlst_results_csv_file(filename, sep='\t'):
    if os.path.exists(filename):
        dataDic = {}
        with open(filename, 'r') as f:
            header = ""
            for n, line in enumerate(f):
                line = line.strip().split(sep)
                if n == 0:
                    header = line
                elif n == 1:
                    dataDic = dict(zip(header, line))

        return dataDic
    else:
        print('\nNo MLST result file {0}\n'.format(filename))
        return ''


def read_species(samplefile, sample_id, sep='\t'):
    if os.path.exists(samplefile):
        with open(samplefile, 'r') as f:
            species = ""
            for n, line in enumerate(f):
                if line.split(sep)[0] == sample_id:
                    species = line.strip().split('\t')[1]
                    # sample_id = line.strip().split('\t')[0]
                    break
        return species
    else:
        print('\nNo sample file {0}\n'.format(samplefile))
        exit(1)


def write_docx(wk_dir, sample_id, species, st, amr_dic, arm_db_name, initial):
    atbs = list(amr_dic.keys())
    atbs.sort()
    fres_dict = {}
    for atb in atbs:
        results = list(amr_dic[atb].keys())
        results.sort()
        for res in results:
            warning = 0
            if 'warning' in amr_dic[atb][res].keys():
                warning = 1
            pc_id = float(amr_dic[atb][res]['id'])
            cov = float(amr_dic[atb][res]['cov'])
            # dep = amr_dic[atb][res]['dep']
            try:
                snp = amr_dic[atb][res]['snp']
                snp = snp.split('[')[0].replace('|', ',')
                # var_only = '1'
            except KeyError:
                snp = ''
                # var_only = '0'

            if pc_id >= 90 and cov >= 80:
                res = res.replace('_', ' ')
                if snp != '' and snp != 'None':
                    snp1 = 'variant résistant de '
                    snp2 = '{0} ({1})'.format(res.strip().split(' ')[0], snp)
                    res = snp1 + snp2
                elif snp == 'None':
                    res = ''
                elif pc_id < 100:
                    res = res + '-like'

                if species in ['klebsiella pneumoniae', 'enterobacter cloacae'] and 'oqx' in res:
                    res = ''

                if 'crrA' in res:
                    res = ''

                if res == 'ampC [Escherichia coli]-like' or res == 'ampC [E. coli]-like' \
                        or res == 'ampC [Escherichia coli]' or res == 'ampC [E. coli]' \
                        or res == 'AmpC [Escherichia coli]' or res == 'AmpC [E. coli]-like' \
                        or res == 'AmpC [Escherichia coli]-like':
                    res = 'AmpC [E. coli]'

                if warning == 1:
                    res = res + '*'

            if res.strip() != '':
                if atb in fres_dict:
                    fres_dict[atb].append(res)
                else:
                    fres_dict[atb] = [res]

    document = Document()
    sections = document.sections
    sections[0].left_margin = Inches(0.5)
    sections[0].right_margin = Inches(0.5)

    ref = document.add_heading('Référence CNR de l\'isolat bacterien : \t', 3)
    ref.add_run(sample_id).italic = True

    spp = document.add_heading('Espèce bactérienne :\t', 3)
    spp.add_run(species[0].upper() + species[1:] + '\n').italic = True

    document.add_heading('Matériel & Méthodes :', 3)
    met_p1 = document.add_paragraph('   ~ Séquençage : ')
    met_p1.add_run('Méthode Illumina (2 x 150pb ou 300pb appariés)\n\n')
    met_p1.add_run('   ~ Analyse ')
    met_p1.add_run('in silico ').italic = True
    met_p1.add_run('des séquences génomiques : Bowtie2, CD-HIT, MUMmer et Samtools\n\n')
    if st != '':
        met_p1.add_run('   ~ Bases de données MLST : ')
        if species == 'escherichia coli':
            met_p1.add_run('http://mlst.warwick.ac.uk/mlst\n\n')
        elif species == 'klebsiella pneumoniae':
            met_p1.add_run('http://bigsdb.pasteur.fr\n\n')
        else:
            met_p1.add_run('https://pubmlst.org/databases\n\n')

    met_p1.add_run(
        '   ~ Bases de données du CNR de la résistance aux antibiotiques : {0} ver.: {1} cat.: {2}\n'
        .format(arm_db_name.split('_')[0], arm_db_name.split('_')[1], arm_db_name.split('_')[2]))

    if st != '':
        document.add_heading('Résultat : Génotypage MLST ', 3)
        document.add_paragraph('   ~ Sequence Type: st-{0}\n'.format(st))

    document.add_heading(
        'Résultat : Déterminants de la résistance aux 3 principales familles d\'antibiotiques (*)', 3)

    frDict = {'Aminoglycoside': "Aminosides", 'Beta-lactam': "Beta-lactamines",
              'Quinolone': "Quinolones", 'Colistin': "Colistine",
              'Sulfonamide': "Sulfamides", 'Trimethoprime': "Triméthoprime",
              'Cycline': "Tétracycline",
              'Aminoglycoside|Fluoroquinolone': "Aminosides et fluoroquinolones"}
    fresults = ['Aminoglycoside', 'Aminoglycoside|Fluoroquinolone', 'Beta-lactam', 'Quinolone',
                'Colistin']  # ,Sulfonamide','Trimethoprime','Cycline',]
    for func in fresults:
        res_txt = ''
        try:
            res = fres_dict[func]
            res.sort()
            res_txt = res_txt + '   ~ {0} : {1}\n'.format(frDict[func], ', '.join(res))
            # document.add_paragraph(res_txt)
        except KeyError:
            res_txt = res_txt + '   ~ {0} :\n'.format(frDict[func])
        document.add_paragraph(res_txt)
        # res_txt = ''

    document.add_paragraph(
        '(*) Contacter le CNR pour d\'autres familles d\'antibiotiques (tél: 04 73 754 920)')
    outfile = os.path.join(os.path.dirname(wk_dir),
                           'CR_CNR_{0}_{1}_{2}.docx'.format(sample_id, datetime.date.today(), initial))
    document.save(outfile)

    print('The main results are summarized in file {0}'.format(outfile))


def main(args):
    wk_dir = args.workdir
    initial = args.initial
    sampleID = os.path.basename(wk_dir)
    samplefile = os.path.join(wk_dir, 'sample.csv')
    species = read_species(samplefile, sampleID, sep='\t')
    try:
        ST_filename = glob.glob(os.path.join(wk_dir, 'mlst_report.tsv'))[0]
        ST_dic = read_mlst_results_csv_file(ST_filename, sep='\t')
        ST = ST_dic['ST']
    except IndexError:
        ST = ''
    arm_filename = glob.glob(os.path.join(wk_dir, 'summary_results_armDB_*.csv'))[0]
    arm_DBname = os.path.splitext(os.path.basename(arm_filename).replace('summary_results_', ''))[0]
    amrDic = read_summary_arm_results_csv_file(arm_filename, sep='\t')
    write_docx(wk_dir, sampleID, species, ST, amrDic, arm_DBname, initial)


def version():
    return "1.0"


def run():
    parser = argparse.ArgumentParser(description='write_docx.py - Version ' + version())
    parser.add_argument('-wd', '--workdir', dest="workdir", default='COL0027', help='working directory')
    parser.add_argument('-in', '--initial', dest="initial", default='RBO', help="Initials of the user")
    parser.add_argument('-V', '--version', action='version', version='write_docx.py-' + version(),
                        help="Prints version number")
    args = parser.parse_args()
    main(args)


if __name__ == '__main__':
    run()
