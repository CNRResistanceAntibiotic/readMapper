#!/usr/bin/python3
# -*- coding: utf-8 -*-
import argparse
import csv
import os
import glob
import datetime

from docx import Document
from docx.shared import Inches


def read_summary_arm_results_csv_file(filename, sep='\t'):
    if os.path.exists(filename):
        data_dict = {}
        with open(filename, 'r') as f:
            header = ""
            for n, line in enumerate(f):
                line = line.strip().split(sep)
                if n == 0:
                    header = line
                elif n == 1:
                    data = line
                    data_dict = dict(zip(header, data))

        atbDic = {}
        for key in data_dict.keys():
            keys = key.split('::')
            atb = keys[0]
            res = keys[1]
            values = data_dict[key].split(',')
            dic = {}
            for value in values:
                key, value = value.split(':')
                dic[key] = value
            try:
                atbDic[atb][res] = dic
            except KeyError:
                atbDic[atb] = {res: dic}

        return atbDic
    else:
        print(f'\nNo armDB result file {filename}\n')
        exit(1)


def read_mlst_results_tsv_file(filename_list):
    data_dic = {}
    for filename in filename_list:
        if os.path.exists(filename):
            with open(filename, 'r') as tsv_file:
                reader = csv.DictReader(tsv_file, delimiter='\t')
                for row in reader:
                    data_dic[row["mlst_name"]] = row
        else:
            print(f'\nNo MLST result file {filename_list}\n')
    return data_dic


def read_species(sample_file, sample_id, sep='\t'):
    if os.path.exists(sample_file):
        with open(sample_file, 'r') as f:
            species = ""
            for n, line in enumerate(f):
                if line.split(sep)[0] == sample_id:
                    species = line.strip().split('\t')[1]
                    # sample_id = line.strip().split('\t')[0]
                    break
        return species
    else:
        print(f'\nNo sample file {sample_file}\n')
        exit(1)


def write_docx(wk_dir, sample_id, species, st_hash, amr_dic, arm_db_name, initial):
    atbs = list(amr_dic.keys())
    atbs.sort()
    f_res_dict = {}
    for atb in atbs:
        results = list(amr_dic[atb].keys())
        results.sort()
        for res in results:
            warning = 0
            if 'warning' in amr_dic[atb][res].keys():
                warning = 1
            pc_id = float(amr_dic[atb][res]['id'])
            cov = float(amr_dic[atb][res]['cov'])
            # dep = amr_dic[atb][res]['dep']
            try:
                snp = amr_dic[atb][res]['snp']
                snp = snp.split('[')[0].replace('|', ',')
                # var_only = '1'
            except KeyError:
                snp = ''
                # var_only = '0'

            if pc_id >= 90 and cov >= 80:
                res = res.replace('_', ' ')
                if snp != '' and snp != 'None':
                    snp1 = 'variant résistant de '
                    snp2 = f'{res.strip().split(" ")[0]} ({snp})'
                    res = snp1 + snp2
                elif snp == 'None':
                    res = ''
                elif pc_id < 100:
                    res = res + '-like'

                if species in ['klebsiella pneumoniae', 'enterobacter cloacae'] and 'oqx' in res:
                    res = ''

                if 'crrA' in res:
                    res = ''

                if res == 'ampC [Escherichia coli]-like' or res == 'ampC [E. coli]-like' \
                        or res == 'ampC [Escherichia coli]' or res == 'ampC [E. coli]' \
                        or res == 'AmpC [Escherichia coli]' or res == 'AmpC [E. coli]-like' \
                        or res == 'AmpC [Escherichia coli]-like':
                    res = 'AmpC [E. coli]'

                if warning == 1:
                    res = res + '*'

            if res.strip() != '':
                if atb in f_res_dict:
                    f_res_dict[atb].append(res)
                else:
                    f_res_dict[atb] = [res]

    document = Document()
    sections = document.sections
    sections[0].left_margin = Inches(0.5)
    sections[0].right_margin = Inches(0.5)

    ref = document.add_heading('Référence CNR de l\'isolat bacterien : \t', 3)
    ref.add_run(sample_id).italic = True

    spp = document.add_heading('Espèce bactérienne :\t', 3)
    spp.add_run(species[0].upper() + species[1:] + '\n').italic = True

    document.add_heading('Matériel & Méthodes :', 3)
    met_p1 = document.add_paragraph('   ~ Séquençage : ')
    met_p1.add_run('Méthode Illumina (2 x 150pb ou 300pb appariés)\n\n')
    met_p1.add_run('   ~ Analyse ')
    met_p1.add_run('in silico ').italic = True
    met_p1.add_run('des séquences génomiques : Ariba, Bowtie2, CD-HIT, MUMmer et Samtools\n\n')

    arm_db_name_split = arm_db_name.split('_')
    met_p1.add_run(f'   ~ Bases de données du CNR de la résistance aux antibiotiques : {arm_db_name_split[0]}'
                   f' ver.: {arm_db_name_split[2]} cat.: {arm_db_name_split[3]}\n')

    if st_hash:
        document.add_heading('Résultat : Génotypage MLST ', 3)
        for schema, st in st_hash.items():
            document.add_paragraph(f'   ~ Sequence Type: ST-{st}. For schema: {schema}\n')

    document.add_heading(
        'Résultat : Déterminants de la résistance aux 3 principales familles d\'antibiotiques (*)', 3)

    fr_dict = {'Aminoglycoside': "Aminosides", 'Beta-lactam': "Beta-lactamines", 'Quinolone': "Quinolones",
               'Colistin': "Colistine", 'Sulfonamide': "Sulfamides", 'Trimethoprime': "Triméthoprime",
               'Cycline': "Tétracycline", 'Aminoglycoside|Fluoroquinolone': "Aminosides et fluoroquinolones"}
    f_results = ['Aminoglycoside', 'Aminoglycoside|Fluoroquinolone', 'Beta-lactam', 'Quinolone', 'Colistin']
    # ,Sulfonamide','Trimethoprime','Cycline',]

    for func in f_results:
        res_txt = ''
        try:
            res = f_res_dict[func]
            res.sort()
            res_txt = res_txt + f'   ~ {fr_dict[func]} : {", ".join(res)}\n'
            # document.add_paragraph(res_txt)
        except KeyError:
            res_txt = res_txt + f'   ~ {fr_dict[func]} :\n'
        document.add_paragraph(res_txt)
        # res_txt = ''

    document.add_paragraph('(*) Contacter le CNR pour d\'autres familles d\'antibiotiques (tél: 04 73 754 920)')
    outfile = os.path.join(os.path.dirname(wk_dir), f'CR_CNR_{sample_id}_{datetime.date.today()}_{initial}.docx')
    document.save(outfile)

    print(f'The main results are summarized in file {outfile}')


def pre_main(args):
    wk_dir = args.workdir
    initial = args.initial
    sample_id = args.sample_id

    # execution main
    main(wk_dir, initial, sample_id)


def main(wk_dir, initial, sample_id):
    wk_dir = os.path.join(wk_dir, sample_id)
    sample_id = os.path.basename(wk_dir)
    sample_file = os.path.join(wk_dir, 'sample.csv')
    species = read_species(sample_file, sample_id, sep='\t')
    st_hash = {}
    try:
        st_filename_list = glob.glob(os.path.join(wk_dir, '*_mlst_report_*.tsv'))
        st_dic = read_mlst_results_tsv_file(st_filename_list)
        for key, value in st_dic.items():
            count_gene = 1
            for col in value:
                if "gene" in col:
                    count_gene += 1
            schema = key
            for count in range(1, count_gene):
                schema = schema + "-" + value[f"gene{count}"]
            st_hash[schema] = value["ST"]
    except IndexError:
        st_list = ''
    arm_filename = glob.glob(os.path.join(wk_dir, 'summary_results_armDB_*.csv'))[0]
    arm_db_name = os.path.splitext(os.path.basename(arm_filename).replace('summary_results_', ''))[0]
    amr_dic = read_summary_arm_results_csv_file(arm_filename, sep='\t')
    write_docx(wk_dir, sample_id, species, st_hash, amr_dic, arm_db_name, initial)


def version():
    return "1.0"


def run():
    parser = argparse.ArgumentParser(description='write_docx.py - Version ' + version())
    parser.add_argument('-wd', '--workdir', dest="workdir", default='COL0027', help='working directory')
    parser.add_argument('-in', '--initial', dest="initial", default='RBO', help="Initials of the user")
    parser.add_argument('-s', '--sampleID', dest="sample_id", default='CNR1979', help='Sample ID')
    parser.add_argument('-V', '--version', action='version', version='write_docx.py-' + version(),
                        help="Prints version number")
    args = parser.parse_args()
    pre_main(args)


if __name__ == '__main__':
    run()
